#!/usr/bin/env python3
"""
Document Text Extraction and Parsing
"""

import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

import PyPDF2
import pdfplumber
from docx import Document
from striprtf.striprtf import rtf_to_text
import chardet

from .config import config

logger = logging.getLogger(__name__)

class DocumentParser:
    """Extracts text and metadata from various document formats"""
    
    def __init__(self):
        self.supported_parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.rtf': self._parse_rtf,
            '.txt': self._parse_text,
            '.md': self._parse_text
        }
    
    def parse_document(self, file_path: Path) -> Dict[str, Any]:
        """Parse document and extract text and metadata"""
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_parsers:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            # Extract text using appropriate parser
            parser_func = self.supported_parsers[file_extension]
            result = parser_func(file_path)
            
            # Add common metadata
            result['metadata'].update({
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension,
                'parsed_at': datetime.now().isoformat(),
                'parser_version': '1.0'
            })
            
            # Validate and clean text
            result['text'] = self._clean_text(result['text'])
            result['word_count'] = len(result['text'].split())
            result['char_count'] = len(result['text'])
            
            # Check text length limits
            if len(result['text']) > config.max_text_length:
                logger.warning(f"Text length ({len(result['text'])}) exceeds limit ({config.max_text_length})")
                result['text'] = result['text'][:config.max_text_length]
                result['truncated'] = True
            else:
                result['truncated'] = False
            
            # Create text chunks for processing
            result['chunks'] = self._create_chunks(result['text'])
            
            logger.info(f"Successfully parsed {file_path.name}: {result['word_count']} words, {len(result['chunks'])} chunks")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse {file_path.name}: {e}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF document"""
        
        text = ""
        metadata = {'pages': 0, 'parser': 'pdf'}
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                # Extract PDF metadata
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                        'producer': pdf.metadata.get('Producer', ''),
                        'creation_date': pdf.metadata.get('CreationDate', ''),
                        'modification_date': pdf.metadata.get('ModDate', '')
                    })
        
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path.name}, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                    
                    # Extract PDF metadata
                    if pdf_reader.metadata:
                        metadata.update({
                            'title': pdf_reader.metadata.get('/Title', ''),
                            'author': pdf_reader.metadata.get('/Author', ''),
                            'subject': pdf_reader.metadata.get('/Subject', ''),
                            'creator': pdf_reader.metadata.get('/Creator', ''),
                            'producer': pdf_reader.metadata.get('/Producer', '')
                        })
                        
                metadata['parser'] = 'pypdf2_fallback'
                
            except Exception as e2:
                logger.error(f"Both PDF parsers failed for {file_path.name}: {e2}")
                raise ValueError(f"Could not parse PDF: {e2}")
        
        return {
            'text': text.strip(),
            'metadata': metadata
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document"""
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            text = "\n\n".join(text_parts)
            
            # Extract metadata
            metadata = {
                'parser': 'docx',
                'paragraphs': len(doc.paragraphs),
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'keywords': doc.core_properties.keywords or '',
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else '',
                'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else '',
                'last_modified_by': doc.core_properties.last_modified_by or ''
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path.name}: {e}")
            raise ValueError(f"Could not parse DOCX: {e}")
    
    def _parse_doc(self, file_path: Path) -> Dict[str, Any]:
        """Parse legacy DOC document"""
        
        # Note: python-docx doesn't support .doc files
        # This is a placeholder for potential future implementation
        # Could use python-docx2txt or antiword if available
        
        logger.warning(f"Legacy DOC format not fully supported: {file_path.name}")
        
        # Try to read as text (will likely fail or produce garbage)
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                
            # Try to detect encoding and extract readable text
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            
            # This is a very basic approach and won't work well for binary DOC files
            text = raw_data.decode(encoding, errors='ignore')
            
            # Clean up binary artifacts
            text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            metadata = {
                'parser': 'doc_basic',
                'encoding': encoding,
                'confidence': detected.get('confidence', 0.0),
                'warning': 'Legacy DOC format - text extraction may be incomplete'
            }
            
            return {
                'text': text.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOC {file_path.name}: {e}")
            raise ValueError(f"Could not parse DOC: {e}")
    
    def _parse_rtf(self, file_path: Path) -> Dict[str, Any]:
        """Parse RTF document"""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            # Extract plain text from RTF
            text = rtf_to_text(rtf_content)
            
            metadata = {
                'parser': 'rtf',
                'original_size': len(rtf_content),
                'extracted_size': len(text)
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to parse RTF {file_path.name}: {e}")
            raise ValueError(f"Could not parse RTF: {e}")
    
    def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text document"""
        
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0.0)
            
            # Read text with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
            
            metadata = {
                'parser': 'text',
                'encoding': encoding,
                'encoding_confidence': confidence,
                'lines': len(text.splitlines())
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to parse text file {file_path.name}: {e}")
            raise ValueError(f"Could not parse text file: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create text chunks for processing"""
        
        if not text:
            return []
        
        chunks = []
        chunk_size = config.chunk_size
        overlap = config.chunk_overlap
        
        # Split text into sentences for better chunking
        sentences = re.split(r'[.!?]+', text)
        
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_size = len(sentence)
            
            # If adding this sentence would exceed chunk size, save current chunk
            if current_size + sentence_size > chunk_size and current_chunk:
                chunks.append({
                    'text': current_chunk.strip(),
                    'size': current_size,
                    'index': len(chunks)
                })
                
                # Start new chunk with overlap
                if overlap > 0 and len(current_chunk) > overlap:
                    overlap_text = current_chunk[-overlap:]
                    current_chunk = overlap_text + " " + sentence
                    current_size = len(overlap_text) + sentence_size + 1
                else:
                    current_chunk = sentence
                    current_size = sentence_size
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += ". " + sentence
                    current_size += sentence_size + 2
                else:
                    current_chunk = sentence
                    current_size = sentence_size
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'size': current_size,
                'index': len(chunks)
            })
        
        logger.info(f"Created {len(chunks)} text chunks")
        return chunks
    
    def extract_metadata_only(self, file_path: Path) -> Dict[str, Any]:
        """Extract only metadata without full text parsing"""
        
        file_extension = file_path.suffix.lower()
        metadata = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_extension': file_extension,
            'extracted_at': datetime.now().isoformat()
        }
        
        try:
            if file_extension == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    metadata['pages'] = len(pdf.pages)
                    if pdf.metadata:
                        metadata.update({
                            'title': pdf.metadata.get('Title', ''),
                            'author': pdf.metadata.get('Author', ''),
                            'subject': pdf.metadata.get('Subject', '')
                        })
            
            elif file_extension == '.docx':
                doc = Document(file_path)
                metadata.update({
                    'paragraphs': len(doc.paragraphs),
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or ''
                })
        
        except Exception as e:
            logger.warning(f"Could not extract metadata from {file_path.name}: {e}")
            metadata['metadata_error'] = str(e)
        
        return metadata

# Global document parser instance
document_parser = DocumentParser()