#!/usr/bin/env python3
"""
File Processing Pipeline for AI Agent Ecosystem
Handles document upload, processing, and organization
"""

import os
import asyncio
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
import hashlib
import mimetypes
from datetime import datetime
import json
import shutil

# Document processing libraries
import PyPDF2
import docx
import markdown
from PIL import Image
import csv

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles file upload, processing, and organization"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        
        # Determine base path based on environment
        if os.name == 'nt' or os.path.exists("X:/"):
            # Windows environment
            self.base_path = Path("X:/AI_Agent_Ecosystem")
        else:
            # Linux environment
            self.base_path = Path("/mnt/llm/AI_Agent_Ecosystem")
        
        # File system structure
        self.userfiles_path = self.base_path / "userfiles"
        self.default_user_path = self.userfiles_path / "default"
        self.uploads_path = self.default_user_path / "uploads"
        self.processed_path = self.default_user_path / "processed"
        self.outputs_path = self.default_user_path / "outputs"
        
        # Supported file types
        self.supported_types = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.json': 'application/json',
            '.csv': 'text/csv',
            '.xml': 'application/xml',
            '.html': 'text/html',
            '.rtf': 'application/rtf'
        }
        
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Create necessary directory structure"""
        try:
            self.uploads_path.mkdir(parents=True, exist_ok=True)
            self.processed_path.mkdir(parents=True, exist_ok=True)
            self.outputs_path.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"File system initialized at: {self.userfiles_path}")
        except Exception as e:
            logger.error(f"Failed to create directories: {e}")
            raise
    
    async def upload_file(self, file_data: bytes, filename: str, 
                         user_id: str = "default") -> Dict[str, Any]:
        """Upload and process a file"""
        try:
            # Validate file type
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.supported_types:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {file_ext}'
                }
            
            # Generate unique filename
            file_hash = hashlib.md5(file_data).hexdigest()[:8]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_filename = f"{timestamp}_{file_hash}_{Path(filename).stem}{file_ext}"
            
            # Save file
            file_path = self.uploads_path / safe_filename
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            # Extract metadata
            metadata = await self._extract_metadata(file_path)
            
            # Process file content
            content = await self._extract_content(file_path)
            
            # Save metadata
            metadata_path = self.processed_path / f"{safe_filename}.meta.json"
            with open(metadata_path, 'w') as f:
                json.dump({
                    'original_filename': filename,
                    'safe_filename': safe_filename,
                    'file_path': str(file_path),
                    'content_preview': content[:500] if content else '',
                    'metadata': metadata,
                    'uploaded_at': datetime.now().isoformat(),
                    'user_id': user_id
                }, f, indent=2)
            
            return {
                'success': True,
                'file_id': safe_filename,
                'file_path': str(file_path),
                'metadata': metadata,
                'content_preview': content[:200] if content else '',
                'size_bytes': len(file_data)
            }
            
        except Exception as e:
            logger.error(f"File upload failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract file metadata"""
        try:
            stat = file_path.stat()
            
            metadata = {
                'size_bytes': stat.st_size,
                'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'file_type': file_path.suffix.lower(),
                'mime_type': mimetypes.guess_type(str(file_path))[0]
            }
            
            # Type-specific metadata
            if file_path.suffix.lower() == '.pdf':
                metadata.update(await self._extract_pdf_metadata(file_path))
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                metadata.update(await self._extract_docx_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {}
    
    async def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'pages': len(reader.pages),
                    'title': reader.metadata.title if reader.metadata else None,
                    'author': reader.metadata.author if reader.metadata else None,
                    'subject': reader.metadata.subject if reader.metadata else None
                }
                
                return metadata
        except Exception as e:
            logger.error(f"PDF metadata extraction failed: {e}")
            return {}
    
    async def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        try:
            doc = docx.Document(file_path)
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'subject': doc.core_properties.subject,
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None
            }
            
            return metadata
        except Exception as e:
            logger.error(f"DOCX metadata extraction failed: {e}")
            return {}
    
    async def _extract_content(self, file_path: Path) -> str:
        """Extract text content from file"""
        try:
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.txt':
                return await self._extract_text_content(file_path)
            elif file_ext == '.md':
                return await self._extract_markdown_content(file_path)
            elif file_ext == '.pdf':
                return await self._extract_pdf_content(file_path)
            elif file_ext in ['.docx', '.doc']:
                return await self._extract_docx_content(file_path)
            elif file_ext == '.json':
                return await self._extract_json_content(file_path)
            elif file_ext == '.csv':
                return await self._extract_csv_content(file_path)
            elif file_ext == '.html':
                return await self._extract_html_content(file_path)
            else:
                return "Content extraction not supported for this file type."
                
        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            return f"Error extracting content: {str(e)}"
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Extract content from text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    async def _extract_markdown_content(self, file_path: Path) -> str:
        """Extract content from markdown file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            # Convert to plain text (remove markdown formatting)
            html = markdown.markdown(content)
            # Simple HTML tag removal
            import re
            text = re.sub('<[^<]+?>', '', html)
            return text
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = []
                
                for page in reader.pages:
                    content.append(page.extract_text())
                
                return '\n'.join(content)
        except Exception as e:
            return f"PDF extraction error: {str(e)}"
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from DOCX file"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            
            return '\n'.join(content)
        except Exception as e:
            return f"DOCX extraction error: {str(e)}"
    
    async def _extract_json_content(self, file_path: Path) -> str:
        """Extract content from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            return f"JSON extraction error: {str(e)}"
    
    async def _extract_csv_content(self, file_path: Path) -> str:
        """Extract content from CSV file"""
        try:
            content = []
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i < 100:  # Limit to first 100 rows
                        content.append(', '.join(row))
                    else:
                        content.append(f"... and {sum(1 for _ in reader)} more rows")
                        break
            
            return '\n'.join(content)
        except Exception as e:
            return f"CSV extraction error: {str(e)}"
    
    async def _extract_html_content(self, file_path: Path) -> str:
        """Extract content from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                # Simple HTML tag removal
                import re
                text = re.sub('<[^<]+?>', '', content)
                return text
        except Exception as e:
            return f"HTML extraction error: {str(e)}"
    
    async def list_files(self, user_id: str = "default") -> List[Dict[str, Any]]:
        """List uploaded files for a user"""
        try:
            files = []
            
            # List files in uploads directory
            for file_path in self.uploads_path.glob('*'):
                if file_path.is_file() and file_path.suffix.lower() in self.supported_types:
                    # Check for metadata file
                    metadata_path = self.processed_path / f"{file_path.name}.meta.json"
                    metadata = {}
                    
                    if metadata_path.exists():
                        with open(metadata_path, 'r') as f:
                            metadata = json.load(f)
                    
                    files.append({
                        'file_id': file_path.name,
                        'original_filename': metadata.get('original_filename', file_path.name),
                        'file_path': str(file_path),
                        'size_bytes': file_path.stat().st_size,
                        'uploaded_at': metadata.get('uploaded_at'),
                        'file_type': file_path.suffix.lower(),
                        'content_preview': metadata.get('content_preview', '')
                    })
            
            return sorted(files, key=lambda x: x.get('uploaded_at', ''), reverse=True)
            
        except Exception as e:
            logger.error(f"Failed to list files: {e}")
            return []
    
    async def get_file_content(self, file_id: str, user_id: str = "default") -> Dict[str, Any]:
        """Get full content of a file"""
        try:
            file_path = self.uploads_path / file_id
            
            if not file_path.exists():
                return {
                    'success': False,
                    'error': 'File not found'
                }
            
            content = await self._extract_content(file_path)
            
            # Get metadata
            metadata_path = self.processed_path / f"{file_id}.meta.json"
            metadata = {}
            if metadata_path.exists():
                with open(metadata_path, 'r') as f:
                    metadata = json.load(f)
            
            return {
                'success': True,
                'file_id': file_id,
                'content': content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to get file content: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def delete_file(self, file_id: str, user_id: str = "default") -> Dict[str, Any]:
        """Delete a file and its metadata"""
        try:
            file_path = self.uploads_path / file_id
            metadata_path = self.processed_path / f"{file_id}.meta.json"
            
            # Delete main file
            if file_path.exists():
                file_path.unlink()
            
            # Delete metadata
            if metadata_path.exists():
                metadata_path.unlink()
            
            return {
                'success': True,
                'message': f'File {file_id} deleted successfully'
            }
            
        except Exception as e:
            logger.error(f"Failed to delete file: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def batch_process_files(self, file_ids: List[str], 
                                workflow: str = "extract_text",
                                user_id: str = "default") -> Dict[str, Any]:
        """Process multiple files in batch"""
        try:
            results = []
            
            for file_id in file_ids:
                result = await self.process_single_file(file_id, workflow, user_id)
                results.append(result)
            
            return {
                'success': True,
                'workflow': workflow,
                'processed_count': len([r for r in results if r.get('success')]),
                'failed_count': len([r for r in results if not r.get('success')]),
                'results': results
            }
            
        except Exception as e:
            logger.error(f"Batch processing failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def process_single_file(self, file_id: str, workflow: str, 
                                user_id: str = "default") -> Dict[str, Any]:
        """Process a single file with specified workflow"""
        try:
            # Get file content
            file_content = await self.get_file_content(file_id, user_id)
            
            if not file_content['success']:
                return file_content
            
            content = file_content['content']
            
            # Apply workflow processing
            if workflow == "extract_text":
                processed_result = {
                    'text_content': content,
                    'word_count': len(content.split()),
                    'character_count': len(content)
                }
            elif workflow == "summarize":
                # Placeholder for summarization - will be replaced with LLM call
                processed_result = {
                    'summary': content[:500] + "..." if len(content) > 500 else content,
                    'original_length': len(content),
                    'summary_ratio': 0.1
                }
            elif workflow == "analyze":
                # Placeholder for analysis - will be replaced with LLM call
                processed_result = {
                    'analysis': f"Document contains {len(content.split())} words and appears to be about: [analysis placeholder]",
                    'key_topics': ["topic1", "topic2", "topic3"],
                    'sentiment': "neutral"
                }
            elif workflow == "categorize":
                # Placeholder for categorization - will be replaced with LLM call
                processed_result = {
                    'category': "general",
                    'confidence': 0.8,
                    'subcategories': ["document", "text"]
                }
            else:
                processed_result = {
                    'error': f'Unknown workflow: {workflow}'
                }
            
            # Save processed result
            output_filename = f"{file_id}_{workflow}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            output_path = self.outputs_path / output_filename
            
            with open(output_path, 'w') as f:
                json.dump({
                    'file_id': file_id,
                    'workflow': workflow,
                    'processed_at': datetime.now().isoformat(),
                    'result': processed_result
                }, f, indent=2)
            
            return {
                'success': True,
                'file_id': file_id,
                'workflow': workflow,
                'result': processed_result,
                'output_file': output_filename
            }
            
        except Exception as e:
            logger.error(f"Single file processing failed: {e}")
            return {
                'success': False,
                'file_id': file_id,
                'error': str(e)
            }
    
    def get_storage_stats(self) -> Dict[str, Any]:
        """Get storage statistics"""
        try:
            def get_dir_size(path: Path) -> int:
                total = 0
                for item in path.rglob('*'):
                    if item.is_file():
                        total += item.stat().st_size
                return total
            
            uploads_size = get_dir_size(self.uploads_path)
            processed_size = get_dir_size(self.processed_path)
            outputs_size = get_dir_size(self.outputs_path)
            
            uploads_count = len(list(self.uploads_path.glob('*')))
            outputs_count = len(list(self.outputs_path.glob('*')))
            
            return {
                'uploads': {
                    'size_bytes': uploads_size,
                    'size_mb': uploads_size / (1024 * 1024),
                    'file_count': uploads_count
                },
                'processed': {
                    'size_bytes': processed_size,
                    'size_mb': processed_size / (1024 * 1024)
                },
                'outputs': {
                    'size_bytes': outputs_size,
                    'size_mb': outputs_size / (1024 * 1024),
                    'file_count': outputs_count
                },
                'total': {
                    'size_bytes': uploads_size + processed_size + outputs_size,
                    'size_mb': (uploads_size + processed_size + outputs_size) / (1024 * 1024)
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to get storage stats: {e}")
            return {}